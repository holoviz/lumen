from io import BytesIO

import param
import pytest

try:
    import lumen.ai  # noqa
except ModuleNotFoundError:
    pytest.skip("lumen.ai could not be imported, skipping tests.", allow_module_level=True)

# python-docx is an optional dependency; skip the whole module without it.
pytest.importorskip("docx")

from docx import Document
from panel.pane import Markdown

from lumen.ai.report import Action, Report, Section


class ChartAction(Action):
    """Emits a pipeline-backed LumenEditor so the story catalog has real views."""

    source = param.Parameter()

    label = param.String(default="Chart")

    async def _execute(self, context, **kwargs):
        from lumen.ai.editors import LumenEditor
        from lumen.pipeline import Pipeline

        pipeline = Pipeline(source=self.source, table='tiny')
        return [LumenEditor(component=pipeline, title=self.label)], {'text': self.label}


class A(Action):

    order = param.List()

    async def _execute(self, context, **kwargs):
        return [Markdown("A done")], {'text': 'A'}


class B(Action):

    order = param.List()

    async def _execute(self, context, **kwargs):
        return [Markdown("B done")], {'text': 'B'}


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def test_docx_add_markdown_headings_bold_bullets():
    from lumen.ai.export import docx_add_markdown

    doc = Document()
    docx_add_markdown(doc, "# Title\n\nSome **bold** text\n\n- item one\n- item two")

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Title" in text
    assert "bold" in text
    assert "item one" in text
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs if p.text == "Title")
    assert any(run.bold for p in doc.paragraphs for run in p.runs if run.text == "bold")


def test_docx_add_table_from_dataframe():
    import pandas as pd

    from lumen.ai.export import docx_add_table

    doc = Document()
    df = pd.DataFrame({"category": ["A", "B"], "value": [1, 2]})
    docx_add_table(doc, df)

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["category", "value"]
    assert len(table.rows) == 3  # header + 2 data rows


def test_docx_add_chart_embeds_png():
    from PIL import Image

    from lumen.ai.export import docx_add_chart

    class FakeEditor:
        export_formats = ("png",)

        def export(self, fmt):
            buf = BytesIO()
            Image.new("RGB", (4, 4), "white").save(buf, "PNG")
            buf.seek(0)
            return buf

    doc = Document()
    assert docx_add_chart(doc, FakeEditor()) is True
    assert len(doc.inline_shapes) == 1


async def test_report_to_docx_headings_and_text():
    report = Report(
        Section(A(order=[]), title='Section A'),
        Section(B(order=[]), title='Section B'),
        title='My Report',
    )
    await report.execute()

    text = _docx_text(report.to_docx())
    assert "My Report" in text
    assert "Section A" in text
    assert "A done" in text
    assert "B done" in text


async def test_report_to_docx_includes_story(llm, tiny_source):
    from lumen.ai.agents.story import Story, StoryBlock

    report = Report(
        Section(ChartAction(source=tiny_source, label='Chart A'), title='Section A'),
        title='My Report', llm=llm,
    )
    await report.execute()
    llm.set_responses([Story(chain_of_thought="c", title="Big Picture", blocks=[
        StoryBlock(prose="Overall narrative."), StoryBlock(view=1),
    ])])
    await report._annotate_report()

    text = _docx_text(report.to_docx())
    assert "Big Picture" in text
    assert "Overall narrative." in text


def test_export_menu_has_word_option():
    report = Report(Section(A(order=[]), title='S'), title='R')
    assert "docx" in [item.get("format") for item in report._export.items]
    assert "docx" in [item.get("format") for item in report._dial.items]


async def test_export_report_docx_returns_bytesio():
    import io

    report = Report(Section(A(order=[]), title='S'), title='My Report')
    await report.execute()

    result = await report._export_report({"format": "docx"})

    assert isinstance(result, io.BytesIO)
    assert report._download.filename == "My Report.docx"
    Document(result)  # the bytes parse as a valid .docx
