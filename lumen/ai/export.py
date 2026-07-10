import base64
import os
import re

from io import BytesIO
from textwrap import dedent
from typing import Any

import nbformat
import yaml

from panel.pane import Markdown
from panel.pane.image import ImageBase
from panel.viewable import Viewable

from ..config import config
from ..pipeline import Pipeline
from ..views import View
from .editors import LumenEditor


def make_md_cell(text: str):
    return nbformat.v4.new_markdown_cell(source=text)

def make_preamble(preamble: str, extensions: list[str] | None = None, title: str | None = None):
    if extensions is None:
        extensions = []
    if 'tabulator' not in extensions:
        extensions = ['tabulator'] + extensions
    exts = ', '.join([repr(ext) for ext in extensions])
    source = (preamble + dedent(
        f"""
        import yaml

        import lumen as lm
        import panel as pn

        pn.extension({exts})
        """
    )).strip()
    imports = nbformat.v4.new_code_cell(source=source)
    if title:
        return [make_md_cell(f'# {title}'), imports]
    return [imports]

def serialize_avatar(avatar: str | bytes | BytesIO | ImageBase, size: int = 45) -> str:
    """
    Process different types of avatar inputs into HTML img tag or text span.

    Args:
        avatar: Avatar source (URL, BytesIO, PIL Image, or text)
        size: Desired width and height of the avatar in pixels

    Returns:
        HTML string representing the avatar
    """
    if isinstance(avatar, ImageBase):
        avatar = avatar.object # type: ignore
    if isinstance(avatar, BytesIO):
        avatar = avatar.getvalue()

    if isinstance(avatar, str):
        if avatar.startswith(('http://', 'https://')):
            return f'<img src="{avatar}" width={size} height={size} alt="avatar" style="border-radius: 50%;"></img>'
        elif os.path.exists(avatar):
            avatar_path = os.path.abspath(avatar)
            return f'<img src="{avatar_path}" width={size} height={size} alt="avatar" style="border-radius: 50%;"></img>'
        else:
            return f'<span class="text-avatar" style="width: {size}px; height: {size}px; display: inline-flex; align-items: center; justify-content: center; background-color: #f0f0f0; border-radius: 50%;">{avatar[:2].upper()}</span>'
    elif isinstance(avatar, bytes):
        img_data = base64.b64encode(avatar).decode()
        return f'<img src="data:image/png;base64,{img_data}" width={size} height={size} alt="avatar" style="border-radius: 50%;"></img>'
    else:
        return f'<span class="text-avatar" style="width: {size}px; height: {size}px; display: inline-flex; align-items: center; justify-content: center; background-color: #f0f0f0; border-radius: 50%;">{str(avatar)[:2].upper()}</span>'

def format_markdown(md: Markdown):
    return [nbformat.v4.new_markdown_cell(source=md.object)]

def format_output(output: Any):
    ext = None
    code = []

    if hasattr(output, 'component') and hasattr(output.component, 'to_spec'):
        component = output.component
    elif hasattr(output, 'to_spec'):
        component = output
    else:
        return None, None

    with config.param.update(serializer='csv'):
        # replace |2- |3- |4-... etc with | for a cleaner look
        try:
            spec = re.sub(r'(\|[-\d]*)', '|', yaml.dump(component.to_spec(), sort_keys=False))
        except (TypeError, yaml.YAMLError):
            return None, None
    read_code = [
        f'yaml_spec = """\n{spec}"""',
        'spec = yaml.safe_load(yaml_spec)',
    ]
    if isinstance(component, Pipeline):
        code.extend([
            *read_code,
            'pipeline = lm.Pipeline.from_spec(spec)',
            'pipeline'
        ])
    elif isinstance(component, View):
        ext = getattr(component, '_extension', None)
        code.extend([
            *read_code,
            'view = lm.View.from_spec(spec)',
            'view'
        ])
    return nbformat.v4.new_code_cell(source='\n'.join(code)), ext

def render_cells(outputs: list[Viewable]) -> tuple[Any, list[str]]:
    cells, extensions = [], []
    for out in outputs:
        if isinstance(out, Markdown):
            cells += format_markdown(out)
        elif isinstance(out, LumenEditor):
            cell, ext = format_output(out)
            if cell is not None:
                cells.append(cell)
            if ext and ext not in extensions:
                extensions.append(ext)
    return cells, extensions

def write_notebook(cells):
    nb = nbformat.v4.new_notebook(cells=cells)
    return nbformat.v4.writes(nb)

def export_notebook(outputs: list[Viewable], preamble: str = ""):
    cells, extensions = render_cells(outputs)
    cells = make_preamble(preamble, extensions=extensions) + cells
    return write_notebook(cells)


def _docx_add_runs(paragraph, text: str):
    """Add text to a paragraph, rendering ``**bold**`` spans as bold runs."""
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def docx_add_markdown(doc, text: str):
    """Render basic markdown (headings, bullets, bold) into a docx document."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            doc.add_heading(stripped.lstrip('# ').strip(), level=min(level, 4))
        elif stripped.startswith(('- ', '* ')):
            _docx_add_runs(doc.add_paragraph(style='List Bullet'), stripped[2:])
        else:
            _docx_add_runs(doc.add_paragraph(), stripped)


def docx_add_table(doc, df, max_rows: int = 50):
    """Render a DataFrame as a native Word table (header + capped rows)."""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    for cell, col in zip(table.rows[0].cells, df.columns, strict=False):
        cell.text = str(col)
    for _, record in df.head(max_rows).iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, record, strict=False):
            cell.text = '' if value is None else str(value)
    if len(df) > max_rows:
        doc.add_paragraph(f'... {len(df) - max_rows} more rows')


def docx_add_chart(doc, editor) -> bool:
    """Embed a chart editor as a PNG image; return True if one was added."""
    from docx.shared import Inches

    if 'png' not in editor.export_formats:
        return False
    try:
        png = editor.export('png')
    except Exception:
        # Best-effort: a chart we cannot render should not abort the document.
        return False
    png.seek(0)
    doc.add_picture(png, width=Inches(6))
    return True
